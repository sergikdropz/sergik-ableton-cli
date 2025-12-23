#!/usr/bin/env python3
"""
Export SERGIK ML Documentation to DOCX

Usage:
    python export_docs.py
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from datetime import datetime


def add_code_block(doc, code: str, language: str = "python"):
    """Add a formatted code block to the document."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def create_documentation():
    """Create the SERGIK ML documentation DOCX."""
    doc = Document()

    # Title
    title = doc.add_heading("SERGIK ML v1.0", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("100% Machine Learning Architecture for Music Production")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Overview",
        "2. Architecture",
        "3. Installation",
        "4. Configuration",
        "5. API Reference",
        "6. Components",
        "7. Database Schema",
        "8. OSC Protocol",
        "9. Training Models",
        "10. Usage Examples",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "SERGIK ML is a complete machine learning system for AI-powered music production. "
        "It provides:\n\n"
        "• Audio feature extraction and embeddings\n"
        "• Preference learning from user ratings\n"
        "• Voice command processing with intent parsing\n"
        "• Sample pack creation pipeline\n"
        "• Real-time Ableton Live integration via OSC\n"
        "• Similarity search for tracks\n"
        "• Adaptive template generation"
    )

    doc.add_heading("Key Features", level=2)
    features = [
        ("Data Ingestion", "Audio files → feature extraction → training datasets"),
        ("Feature Extraction", "BPM, energy, brightness, LUFS, harmonic/percussive ratio"),
        ("Embeddings", "Text-to-style embeddings for prompt-based generation"),
        ("Models", "Intent parsing, preference learning, template proposal, reranking"),
        ("Serving", "FastAPI REST API with voice endpoint"),
        ("Connectors", "Ableton OSC, PostgreSQL (optional), cloud storage"),
    ]
    for name, desc in features:
        p = doc.add_paragraph()
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # 2. Architecture
    doc.add_heading("2. Architecture", level=1)

    doc.add_heading("System Diagram", level=2)
    arch_diagram = """
┌─────────────────────────────────────────────────────────────────┐
│                        SERGIK ML v1.0                           │
├─────────────────────────────────────────────────────────────────┤
│                                                                 │
│  ┌─────────────┐    ┌─────────────┐    ┌─────────────┐         │
│  │   FastAPI   │    │   Voice     │    │   Ableton   │         │
│  │   /action   │◄──►│   Pipeline  │◄──►│   OSC       │         │
│  └──────┬──────┘    └──────┬──────┘    └──────┬──────┘         │
│         │                  │                  │                 │
│  ┌──────▼──────────────────▼──────────────────▼──────┐         │
│  │                  Action Dispatcher                │         │
│  └──────┬──────────────────┬──────────────────┬──────┘         │
│         │                  │                  │                 │
│  ┌──────▼──────┐    ┌──────▼──────┐    ┌──────▼──────┐         │
│  │    Pack     │    │  Similarity │    │    Live     │         │
│  │  Pipeline   │    │   Search    │    │  Commands   │         │
│  └──────┬──────┘    └──────┬──────┘    └─────────────┘         │
│         │                  │                                    │
│  ┌──────▼──────────────────▼──────────────────────────┐        │
│  │                   SQL Store                        │        │
│  │  (music_intelligence, pack_manifests, action_log)  │        │
│  └────────────────────────────────────────────────────┘        │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, arch_diagram, "text")

    doc.add_heading("Package Structure", level=2)
    structure = """
sergik_ml/
├── __init__.py           # Package metadata
├── config.py             # Configuration from environment
├── schemas.py            # Pydantic models
├── stores/
│   ├── sql_store.py      # SQLAlchemy database
│   └── vector_store.py   # Similarity search
├── connectors/
│   └── ableton_osc.py    # OSC client
├── features/
│   ├── audio_features.py # Audio analysis
│   └── text_embeddings.py# Text-to-style
├── models/
│   ├── preference.py     # Rating-based learning
│   ├── template.py       # Pack templates
│   ├── intent.py         # Voice command parsing
│   └── rerank.py         # Similarity reranking
├── policies/
│   └── action_policy.py  # Command validation
├── pipelines/
│   ├── pack_pipeline.py  # Sample pack creation
│   └── voice_pipeline.py # STT/TTS processing
├── serving/
│   └── api.py            # FastAPI application
└── train/
    └── train_preference.py# Model training
"""
    add_code_block(doc, structure, "text")

    doc.add_page_break()

    # 3. Installation
    doc.add_heading("3. Installation", level=1)

    doc.add_heading("Requirements", level=2)
    doc.add_paragraph("Python 3.9+ required.")

    doc.add_heading("Install Dependencies", level=2)
    add_code_block(doc, """pip install -r requirements.txt""", "bash")

    doc.add_heading("Core Dependencies", level=2)
    deps = [
        "fastapi, uvicorn - Web server",
        "sqlalchemy - Database ORM",
        "numpy, scipy, pandas - Data processing",
        "librosa, soundfile, pyloudnorm - Audio analysis",
        "python-osc - Ableton communication",
        "mido - MIDI file handling",
        "click - CLI framework",
    ]
    for dep in deps:
        doc.add_paragraph(f"• {dep}")

    doc.add_page_break()

    # 4. Configuration
    doc.add_heading("4. Configuration", level=1)
    doc.add_paragraph("Configuration via environment variables:")

    config_table = doc.add_table(rows=1, cols=3)
    config_table.style = "Table Grid"
    hdr_cells = config_table.rows[0].cells
    hdr_cells[0].text = "Variable"
    hdr_cells[1].text = "Default"
    hdr_cells[2].text = "Description"

    configs = [
        ("SERGIK_HOST", "127.0.0.1", "Server bind host"),
        ("SERGIK_PORT", "8000", "Server bind port"),
        ("SERGIK_DB_URL", "sqlite:///sergik_ml.db", "Database URL"),
        ("SERGIK_ARTIFACT_DIR", "artifacts", "Model storage"),
        ("SERGIK_ABLETON_OSC_HOST", "127.0.0.1", "Ableton OSC host"),
        ("SERGIK_ABLETON_OSC_PORT", "9000", "Ableton OSC port"),
        ("SERGIK_USE_OPENAI_VOICE", "0", "Enable OpenAI voice"),
        ("OPENAI_API_KEY", "", "OpenAI API key"),
    ]

    for var, default, desc in configs:
        row_cells = config_table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = default
        row_cells[2].text = desc

    doc.add_page_break()

    # 5. API Reference
    doc.add_heading("5. API Reference", level=1)

    doc.add_heading("Endpoints", level=2)

    endpoints = [
        ("GET /health", "Health check", '{"status": "ok", "service": "sergik-ml"}'),
        ("POST /action", "Execute command", '{"cmd": "pack.create", "args": {...}}'),
        ("POST /voice", "Process voice (upload WAV)", "Returns intent and TTS path"),
        ("GET /tracks", "List tracks", "Query: limit, rated_only"),
        ("GET /tracks/{id}", "Get track details", "Returns track metadata"),
        ("GET /similar/{id}", "Find similar tracks", "Query: k, style"),
    ]

    for endpoint, desc, notes in endpoints:
        doc.add_heading(endpoint, level=3)
        doc.add_paragraph(desc)
        p = doc.add_paragraph()
        p.add_run("Notes: ").bold = True
        p.add_run(notes)

    doc.add_heading("Commands", level=2)

    commands = [
        ("pack.create", "Create sample pack", "stems_dir, tempo, length_bars, auto_zip"),
        ("pack.rate", "Rate a track", "track_id, rating (1-5)"),
        ("pack.similar", "Find similar tracks", "track_id, k"),
        ("live.set_tempo", "Set Ableton tempo", "tempo"),
        ("live.play", "Start playback", "-"),
        ("live.stop", "Stop playback", "-"),
        ("gen.generate_loop", "Generate audio", "prompt, style, bpm, bars"),
        ("voice.tts", "Text-to-speech", "text"),
    ]

    cmd_table = doc.add_table(rows=1, cols=3)
    cmd_table.style = "Table Grid"
    hdr = cmd_table.rows[0].cells
    hdr[0].text = "Command"
    hdr[1].text = "Description"
    hdr[2].text = "Arguments"

    for cmd, desc, args in commands:
        row = cmd_table.add_row().cells
        row[0].text = cmd
        row[1].text = desc
        row[2].text = args

    doc.add_page_break()

    # 6. Components
    doc.add_heading("6. Components", level=1)

    doc.add_heading("Audio Feature Extraction", level=2)
    doc.add_paragraph("Extracts ML-ready features from audio files:")
    features_list = [
        "BPM/Tempo (librosa beat tracking)",
        "Energy (RMS amplitude)",
        "Brightness (spectral centroid)",
        "LUFS loudness (pyloudnorm)",
        "Harmonic/Percussive ratio (librosa HPSS)",
        "Key estimation (chroma analysis)",
        "MFCCs for timbral fingerprint",
    ]
    for f in features_list:
        doc.add_paragraph(f"• {f}")

    doc.add_heading("Preference Model", level=2)
    doc.add_paragraph(
        "Ridge regression model trained from user ratings. "
        "Learns weights for audio features to predict user preference."
    )
    add_code_block(doc, """
# Training
from sergik_ml.models.preference import PreferenceModel
model = PreferenceModel()
model.fit(X, y, l2=1e-2)
model.save("preference_model.pkl")

# Prediction
score = model.predict(feature_vector)
""", "python")

    doc.add_heading("Intent Model", level=2)
    doc.add_paragraph(
        "Rule-based intent parser for voice commands. "
        "Matches patterns to extract {cmd, args, tts}."
    )
    doc.add_paragraph("Supported intents:")
    intents = [
        '"make a 4-bar sample pack" → pack.create',
        '"find similar loops" → pack.similar',
        '"set tempo 125" → live.set_tempo',
        '"rate 5 stars" → pack.rate',
    ]
    for intent in intents:
        doc.add_paragraph(f"• {intent}")

    doc.add_page_break()

    # 7. Database Schema
    doc.add_heading("7. Database Schema", level=1)

    doc.add_heading("music_intelligence", level=2)
    doc.add_paragraph("Track features and metadata for ML training.")
    mi_cols = [
        ("track_id", "VARCHAR(255)", "Primary key"),
        ("bpm", "FLOAT", "Tempo"),
        ("key", "VARCHAR(10)", "Musical key"),
        ("energy", "FLOAT", "RMS energy"),
        ("brightness", "FLOAT", "Spectral centroid"),
        ("lufs", "FLOAT", "Integrated loudness"),
        ("harmonic_ratio", "FLOAT", "Harmonic content ratio"),
        ("percussive_ratio", "FLOAT", "Percussive content ratio"),
        ("rating", "FLOAT", "User rating (1-5)"),
        ("source_pack", "VARCHAR(255)", "Source pack name"),
        ("style_source", "VARCHAR(50)", "Style origin"),
        ("updated_at", "DATETIME", "Last update"),
    ]
    mi_table = doc.add_table(rows=1, cols=3)
    mi_table.style = "Table Grid"
    hdr = mi_table.rows[0].cells
    hdr[0].text = "Column"
    hdr[1].text = "Type"
    hdr[2].text = "Description"
    for col, typ, desc in mi_cols:
        row = mi_table.add_row().cells
        row[0].text = col
        row[1].text = typ
        row[2].text = desc

    doc.add_heading("action_log", level=2)
    doc.add_paragraph("All command executions for ML training data.")

    doc.add_heading("pack_manifests", level=2)
    doc.add_paragraph("Sample pack exports with metadata.")

    doc.add_heading("emotion_intelligence", level=2)
    doc.add_paragraph("Preference/emotion signals for training.")

    doc.add_page_break()

    # 8. OSC Protocol
    doc.add_heading("8. OSC Protocol", level=1)
    doc.add_paragraph("Communication between SERGIK ML and Max for Live.")

    doc.add_heading("ML Service → Ableton", level=2)
    osc_out = [
        ("/scp/status", '{"text": "..."}', "Status message"),
        ("/scp/similar_results", '{"track_id": "...", "similar": [...]}', "Search results"),
        ("/scp/tts_ready", '{"path": "..."}', "TTS audio ready"),
        ("/scp/pack_ready", '{"pack_id": "...", "export_dir": "..."}', "Pack created"),
        ("/scp/error", '{"error": "..."}', "Error message"),
        ("/scp/set_tempo", '{"tempo": 125}', "Set tempo"),
    ]
    osc_table = doc.add_table(rows=1, cols=3)
    osc_table.style = "Table Grid"
    hdr = osc_table.rows[0].cells
    hdr[0].text = "Address"
    hdr[1].text = "Payload"
    hdr[2].text = "Description"
    for addr, payload, desc in osc_out:
        row = osc_table.add_row().cells
        row[0].text = addr
        row[1].text = payload
        row[2].text = desc

    doc.add_heading("Max for Live Setup", level=2)
    add_code_block(doc, """
[udpreceive 9000]
    |
[route /scp/status /scp/similar_results /scp/tts_ready]
    |
[fromsymbol] -> [dict.deserialize] -> [js handler.js]
""", "text")

    doc.add_page_break()

    # 9. Training Models
    doc.add_heading("9. Training Models", level=1)

    doc.add_heading("Train Preference Model", level=2)
    add_code_block(doc, """
# Requires rated tracks in database
python -m sergik_ml.train.train_preference

# Output: artifacts/models/preference_model.pkl
""", "bash")

    doc.add_heading("Training Data Collection", level=2)
    doc.add_paragraph("The system automatically collects training data:")
    data_sources = [
        "action_log: All command executions",
        "emotion_intelligence: Rating-derived signals",
        "music_intelligence: Track features",
    ]
    for src in data_sources:
        doc.add_paragraph(f"• {src}")

    doc.add_page_break()

    # 10. Usage Examples
    doc.add_heading("10. Usage Examples", level=1)

    doc.add_heading("Start Server", level=2)
    add_code_block(doc, """python run_server.py""", "bash")

    doc.add_heading("Create Sample Pack", level=2)
    add_code_block(doc, """
curl -X POST http://localhost:8000/action \\
  -H "Content-Type: application/json" \\
  -d '{
    "cmd": "pack.create",
    "args": {
      "stems_dir": "uploads",
      "tempo": 125,
      "length_bars": 4,
      "auto_zip": true
    }
  }'
""", "bash")

    doc.add_heading("Rate Track", level=2)
    add_code_block(doc, """
curl -X POST http://localhost:8000/action \\
  -H "Content-Type: application/json" \\
  -d '{
    "cmd": "pack.rate",
    "args": {"track_id": "loop_01_4bar", "rating": 5}
  }'
""", "bash")

    doc.add_heading("Find Similar", level=2)
    add_code_block(doc, """
curl http://localhost:8000/similar/loop_01_4bar?k=10
""", "bash")

    doc.add_heading("Voice Command", level=2)
    add_code_block(doc, """
curl -X POST http://localhost:8000/voice \\
  -F "file=@voice_recording.wav"
""", "bash")

    # Save document
    output_path = Path("/Users/machd/sergik_custom_gpt/SERGIK_ML_Documentation.docx")
    doc.save(str(output_path))
    print(f"Documentation saved to: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    create_documentation()
